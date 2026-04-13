"""
Routes API pour le module Documentations - Pôles de Service et Documents
"""
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import HTMLResponse
from fastapi.responses import StreamingResponse
from typing import List, Optional
from datetime import datetime, timezone
from pathlib import Path
import uuid
import logging
import mimetypes
import json
import os
from io import BytesIO

from models import (
    PoleDeService,
    PoleDeServiceCreate,
    PoleDeServiceUpdate,
    Document,
    DocumentCreate,
    DocumentUpdate,
    BonDeTravail,
    BonDeTravailCreate,
    DocumentType,
    ServicePole,
    ActionType,
    EntityType,
    SuccessResponse
)
from dependencies import get_current_user, get_current_admin_user, get_current_user_optional
from audit_service import AuditService
from auth import decode_access_token
from bon_travail_template_final import generate_bon_travail_html

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documentations", tags=["documentations"])

# Variables globales (injectées depuis server.py)
db = None
audit_service = None
realtime_manager = None

def init_documentations_routes(database, audit_svc, realtime_mgr=None):
    """Initialise les routes avec la connexion DB, audit service et realtime manager"""
    global db, audit_service, realtime_manager
    db = database
    audit_service = audit_svc
    realtime_manager = realtime_mgr


# ==================== PÔLES DE SERVICE ====================

@router.get("/poles", response_model=List[dict])
async def get_poles(current_user: dict = Depends(get_current_user)):
    """Récupérer tous les pôles de service avec leurs documents et bons de travail"""
    try:
        poles = await db.poles_service.find().to_list(length=None)
        
        # Pour chaque pôle, récupérer les documents et bons associés
        for pole in poles:
            if "_id" in pole:
                del pole["_id"]
            
            # Récupérer les documents associés
            documents = await db.documents.find({"pole_id": pole["id"]}).to_list(length=None)
            for doc in documents:
                if "_id" in doc:
                    del doc["_id"]
            
            # Récupérer les bons de travail associés
            bons_travail = await db.bons_travail.find({"pole_id": pole["id"]}).to_list(length=None)
            for bon in bons_travail:
                if "_id" in bon:
                    del bon["_id"]
            
            # Ajouter les documents et bons au pôle
            pole["documents"] = documents
            pole["bons_travail"] = bons_travail
        
        return poles
    except Exception as e:
        logger.error(f"Erreur récupération pôles: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/poles/{pole_id}")
async def get_pole(pole_id: str, current_user: dict = Depends(get_current_user)):
    """Récupérer un pôle spécifique avec ses documents et bons de travail"""
    try:
        pole = await db.poles_service.find_one({"id": pole_id})
        if not pole:
            raise HTTPException(status_code=404, detail="Pôle non trouvé")
        if "_id" in pole:
            del pole["_id"]
        
        # Récupérer les documents associés au pôle
        documents = await db.documents.find({"pole_id": pole_id}).to_list(length=None)
        for doc in documents:
            if "_id" in doc:
                del doc["_id"]
        
        # Récupérer les bons de travail associés au pôle
        bons_travail = await db.bons_travail.find({"pole_id": pole_id}).to_list(length=None)
        for bon in bons_travail:
            if "_id" in bon:
                del bon["_id"]
        
        # Ajouter les documents et bons au pôle
        pole["documents"] = documents
        pole["bons_travail"] = bons_travail
        
        return pole
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération pôle {pole_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/poles")
async def create_pole(
    pole_data: PoleDeServiceCreate,
    current_user: dict = Depends(get_current_user)
):
    """Créer un nouveau pôle de service"""
    try:
        pole = PoleDeService(
            **pole_data.model_dump(),
            created_by=current_user.get("id")
        )
        
        pole_dict = pole.model_dump()
        await db.poles_service.insert_one(pole_dict)
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.CREATE,
            entity_type=EntityType.SETTINGS,
            entity_id=pole.id,
            entity_name=f"Pôle: {pole.nom}"
        )
        
        if "_id" in pole_dict:
            del pole_dict["_id"]
        
        # Broadcast WebSocket pour la synchronisation temps réel
        if realtime_manager:
            await realtime_manager.emit_event(
                "documentations",
                "created",
                pole_dict,
                user_id=current_user["id"]
            )
        
        return pole_dict
    except Exception as e:
        logger.error(f"Erreur création pôle: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/poles/{pole_id}")
async def update_pole(
    pole_id: str,
    pole_update: PoleDeServiceUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Mettre à jour un pôle de service"""
    try:
        existing = await db.poles_service.find_one({"id": pole_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Pôle non trouvé")
        
        update_data = {
            k: v for k, v in pole_update.model_dump(exclude_unset=True).items()
            if v is not None
        }
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        
        await db.poles_service.update_one(
            {"id": pole_id},
            {"$set": update_data}
        )
        
        updated_pole = await db.poles_service.find_one({"id": pole_id})
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.UPDATE,
            entity_type=EntityType.SETTINGS,
            entity_id=pole_id,
            entity_name=f"Pôle: {existing.get('nom')}"
        )
        
        if "_id" in updated_pole:
            del updated_pole["_id"]
        
        # Broadcast WebSocket pour la synchronisation temps réel
        if realtime_manager:
            await realtime_manager.emit_event(
                "documentations",
                "updated",
                updated_pole,
                user_id=current_user["id"]
            )
        
        return updated_pole
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur mise à jour pôle {pole_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/poles/{pole_id}", response_model=SuccessResponse)
async def delete_pole(
    pole_id: str,
    current_user: dict = Depends(get_current_admin_user)
):
    """Supprimer un pôle de service (Admin uniquement)"""
    try:
        pole = await db.poles_service.find_one({"id": pole_id})
        if not pole:
            raise HTTPException(status_code=404, detail="Pôle non trouvé")
        
        # Vérifier s'il y a des documents liés
        docs_count = await db.documents.count_documents({"pole_id": pole_id})
        if docs_count > 0:
            raise HTTPException(
                status_code=400,
                detail=f"Impossible de supprimer: {docs_count} document(s) lié(s) à ce pôle"
            )
        
        await db.poles_service.delete_one({"id": pole_id})
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.DELETE,
            entity_type=EntityType.SETTINGS,
            entity_id=pole_id,
            entity_name=f"Pôle: {pole.get('nom')}"
        )
        
        # Broadcast WebSocket pour la synchronisation temps réel
        if realtime_manager:
            await realtime_manager.emit_event(
                "documentations",
                "deleted",
                {"id": pole_id, "nom": pole.get('nom')},
                user_id=current_user["id"]
            )
        
        return {"success": True, "message": "Pôle supprimé"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur suppression pôle {pole_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== DOCUMENTS ====================

@router.get("/documents", response_model=List[dict])
async def get_documents(
    pole_id: Optional[str] = None,
    type_document: Optional[str] = None,
    statut: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer tous les documents avec filtres"""
    try:
        query = {}
        if pole_id:
            query["pole_id"] = pole_id
        if type_document:
            query["type_document"] = type_document
        if statut:
            query["statut"] = statut
        
        documents = await db.documents.find(query).to_list(length=None)
        for doc in documents:
            if "_id" in doc:
                del doc["_id"]
        return documents
    except Exception as e:
        logger.error(f"Erreur récupération documents: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/documents/{document_id}")
async def get_document(
    document_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer un document spécifique"""
    try:
        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")
        if "_id" in doc:
            del doc["_id"]
        return doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération document {document_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/documents")
async def create_document(
    doc_data: DocumentCreate,
    current_user: dict = Depends(get_current_user)
):
    """Créer un nouveau document"""
    try:
        # Vérifier que le pôle existe
        pole = await db.poles_service.find_one({"id": doc_data.pole_id})
        if not pole:
            raise HTTPException(status_code=404, detail="Pôle non trouvé")
        
        doc = Document(
            **doc_data.model_dump(),
            created_by=current_user.get("id"),
            updated_by=current_user.get("id")
        )
        
        doc_dict = doc.model_dump()
        await db.documents.insert_one(doc_dict)
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.CREATE,
            entity_type=EntityType.SETTINGS,
            entity_id=doc.id,
            entity_name=f"Document: {doc.titre}"
        )
        
        if "_id" in doc_dict:
            del doc_dict["_id"]
        
        # Broadcast WebSocket pour la synchronisation temps réel
        if realtime_manager:
            await realtime_manager.emit_event(
                "documentations",
                "created",
                doc_dict,
                user_id=current_user["id"]
            )
        
        return doc_dict
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur création document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/documents/{document_id}")
async def update_document(
    document_id: str,
    doc_update: DocumentUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Mettre à jour un document"""
    try:
        existing = await db.documents.find_one({"id": document_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Document non trouvé")
        
        update_data = {
            k: v for k, v in doc_update.model_dump(exclude_unset=True).items()
            if v is not None
        }
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()
        update_data["updated_by"] = current_user.get("id")
        
        await db.documents.update_one(
            {"id": document_id},
            {"$set": update_data}
        )
        
        updated_doc = await db.documents.find_one({"id": document_id})
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.UPDATE,
            entity_type=EntityType.SETTINGS,
            entity_id=document_id,
            entity_name=f"Document: {existing.get('titre')}"
        )
        
        if "_id" in updated_doc:
            del updated_doc["_id"]
        
        # Broadcast WebSocket pour la synchronisation temps réel
        if realtime_manager:
            await realtime_manager.emit_event(
                "documentations",
                "updated",
                updated_doc,
                user_id=current_user["id"]
            )
        
        return updated_doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur mise à jour document {document_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/documents/{document_id}", response_model=SuccessResponse)
async def delete_document(
    document_id: str,
    current_user: dict = Depends(get_current_admin_user)
):
    """Supprimer un document"""
    try:
        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")
        
        # Supprimer le fichier physique si c'est une pièce jointe
        if doc.get("fichier_url"):
            try:
                file_path = Path(f"/app{doc['fichier_url']}")
                if file_path.exists():
                    file_path.unlink()
            except Exception as e:
                logger.warning(f"Impossible de supprimer le fichier: {e}")
        
        await db.documents.delete_one({"id": document_id})
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.DELETE,
            entity_type=EntityType.SETTINGS,
            entity_id=document_id,
            entity_name=f"Document: {doc.get('titre')}"
        )
        
        return {"success": True, "message": "Document supprimé"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur suppression document {document_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== UPLOAD FICHIERS ====================

@router.post("/documents/{document_id}/upload")
async def upload_document_file(
    document_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Upload un fichier pour un document"""
    try:
        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")
        
        # Créer le répertoire uploads/documents si nécessaire
        upload_dir = Path("uploads/documents")
        upload_dir.mkdir(parents=True, exist_ok=True)
        
        # Générer un nom de fichier unique
        file_ext = Path(file.filename).suffix
        unique_filename = f"{document_id}_{uuid.uuid4()}{file_ext}"
        file_path = upload_dir / unique_filename
        
        # Sauvegarder le fichier
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Déterminer le type MIME
        mime_type, _ = mimetypes.guess_type(file.filename)
        
        # Mettre à jour le document avec les infos du fichier
        file_url = f"/uploads/documents/{unique_filename}"
        await db.documents.update_one(
            {"id": document_id},
            {
                "$set": {
                    "fichier_url": file_url,
                    "fichier_nom": file.filename,
                    "fichier_type": mime_type or "application/octet-stream",
                    "fichier_taille": len(content),
                    "updated_at": datetime.now(timezone.utc).isoformat(),
                    "updated_by": current_user.get("id")
                }
            }
        )
        
        return {
            "success": True,
            "file_url": file_url,
            "file_name": file.filename,
            "file_size": len(content),
            "file_type": mime_type
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur upload fichier: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/upload-files")
async def upload_files_direct(
    files: List[UploadFile] = File(...),
    pole_id: str = Form(...),
    folder_id: Optional[str] = Form(None),
    current_user: dict = Depends(get_current_user)
):
    """Upload direct de fichiers dans un dossier - crée automatiquement les documents"""
    try:
        upload_dir = Path("uploads/documents")
        upload_dir.mkdir(parents=True, exist_ok=True)

        created_docs = []
        for file in files:
            doc_id = str(uuid.uuid4())
            file_ext = Path(file.filename).suffix
            unique_filename = f"{doc_id}_{uuid.uuid4()}{file_ext}"
            file_path = upload_dir / unique_filename

            content = await file.read()
            with open(file_path, "wb") as f:
                f.write(content)

            mime_type, _ = mimetypes.guess_type(file.filename)
            file_url = f"/uploads/documents/{unique_filename}"

            # Créer le document dans la base
            doc_dict = {
                "id": doc_id,
                "pole_id": pole_id,
                "folder_id": folder_id if folder_id and folder_id != "null" else None,
                "titre": Path(file.filename).stem,
                "type_document": "fichier",
                "fichier_url": file_url,
                "fichier_nom": file.filename,
                "fichier_type": mime_type or "application/octet-stream",
                "fichier_taille": len(content),
                "created_at": datetime.now(timezone.utc).isoformat(),
                "updated_at": datetime.now(timezone.utc).isoformat(),
                "created_by": current_user.get("id"),
                "updated_by": current_user.get("id"),
                "hidden_for_external": False,
                "hidden_for_users": False
            }
            await db.documents.insert_one(doc_dict)
            doc_dict.pop("_id", None)
            created_docs.append(doc_dict)

            # WebSocket
            if realtime_manager:
                await realtime_manager.emit_event("documentations", "created", doc_dict, user_id=current_user["id"])

            # Audit
            await audit_service.log_action(
                user_id=current_user["id"],
                user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
                user_email=current_user.get("email", ""),
                action=ActionType.CREATE,
                entity_type=EntityType.DOCUMENTATION,
                entity_id=doc_id,
                entity_name=file.filename,
                details=f"a ajouté le fichier \"{file.filename}\""
            )

        return {"success": True, "count": len(created_docs), "documents": created_docs}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur upload fichiers: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/documents/{document_id}/view")
async def view_document_file(
    document_id: str,
    token: str = None,
    current_user: dict = Depends(get_current_user_optional)
):
    """Visualiser le fichier d'un document dans le navigateur (inline)"""
    try:
        # Si pas d'utilisateur via Bearer token, vérifier le token en query param
        if not current_user and token:
            # Vérifier le token passé en paramètre
            payload = decode_access_token(token)
            if payload is None:
                raise HTTPException(status_code=401, detail="Token invalide ou expiré")
        elif not current_user and not token:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")
        
        if not doc.get("fichier_url"):
            raise HTTPException(status_code=404, detail="Aucun fichier associé")
        
        # Le fichier_url commence par /uploads/documents/
        # Le fichier réel est dans /app/backend/uploads/documents/
        file_path = Path(f"/app/backend{doc['fichier_url']}")
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Fichier non trouvé sur le serveur")
        
        # Lire le fichier
        with open(file_path, "rb") as f:
            content = f.read()
        
        # Utiliser inline pour permettre la visualisation dans le navigateur
        return StreamingResponse(
            BytesIO(content),
            media_type=doc.get("fichier_type", "application/octet-stream"),
            headers={
                "Content-Disposition": f"inline; filename={doc.get('fichier_nom', 'document')}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur visualisation fichier: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/documents/{document_id}/download")
async def download_document_file(
    document_id: str,
    token: str = None,
    current_user: dict = Depends(get_current_user_optional)
):
    """Télécharger le fichier d'un document (force le téléchargement)"""
    try:
        if not current_user and token:
            payload = decode_access_token(token)
            if payload is None:
                raise HTTPException(status_code=401, detail="Token invalide ou expiré")
        elif not current_user and not token:
            raise HTTPException(status_code=401, detail="Not authenticated")

        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")
        
        if not doc.get("fichier_url"):
            raise HTTPException(status_code=404, detail="Aucun fichier associé")
        
        file_path = Path(f"/app/backend{doc['fichier_url']}")
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Fichier non trouvé sur le serveur")
        
        with open(file_path, "rb") as f:
            content = f.read()
        
        filename = doc.get('fichier_nom', 'document')
        return StreamingResponse(
            BytesIO(content),
            media_type="application/octet-stream",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(content))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur téléchargement fichier: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== BON DE TRAVAIL ====================

@router.get("/bons-travail", response_model=List[dict])
async def get_bons_travail(
    pole_id: Optional[str] = None,
    statut: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer tous les bons de travail"""
    try:
        query = {}
        if pole_id:
            query["pole_id"] = pole_id
        if statut:
            query["statut"] = statut
        
        bons = await db.bons_travail.find(query).to_list(length=None)
        for bon in bons:
            if "_id" in bon:
                del bon["_id"]
        return bons
    except Exception as e:
        logger.error(f"Erreur récupération bons de travail: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/bons-travail/{bon_id}")
async def get_bon_travail(
    bon_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer un bon de travail spécifique"""
    try:
        bon = await db.bons_travail.find_one({"id": bon_id})
        if not bon:
            raise HTTPException(status_code=404, detail="Bon de travail non trouvé")
        if "_id" in bon:
            del bon["_id"]
        return bon
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération bon {bon_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/bons-travail")
async def create_bon_travail(
    bon_data: BonDeTravailCreate,
    current_user: dict = Depends(get_current_user)
):
    """Créer un nouveau bon de travail"""
    try:
        bon = BonDeTravail(
            **bon_data.model_dump(),
            created_by=current_user.get("id")
        )
        
        bon_dict = bon.model_dump()
        await db.bons_travail.insert_one(bon_dict)
        
        # Audit
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user['prenom']} {current_user['nom']}",
            user_email=current_user["email"],
            action=ActionType.CREATE,
            entity_type=EntityType.SETTINGS,
            entity_id=bon.id,
            entity_name=f"Bon de travail: {bon.localisation_ligne}"
        )
        
        if "_id" in bon_dict:
            del bon_dict["_id"]
        
        return bon_dict
    except Exception as e:
        logger.error(f"Erreur création bon de travail: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/bons-travail/{bon_id}")
async def update_bon_travail(
    bon_id: str,
    bon_update: dict,
    current_user: dict = Depends(get_current_user)
):
    """Mettre à jour un bon de travail - Permissions : admin ou créateur uniquement"""
    try:
        existing = await db.bons_travail.find_one({"id": bon_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Bon de travail non trouvé")
        
        # Vérifier les permissions : admin ou créateur
        if current_user.get("role") != "ADMIN" and existing.get("created_by") != current_user.get("id"):
            raise HTTPException(
                status_code=403, 
                detail="Vous n'avez pas la permission de modifier ce bon de travail"
            )
        
        bon_update["updated_at"] = datetime.now(timezone.utc).isoformat()
        
        await db.bons_travail.update_one(
            {"id": bon_id},
            {"$set": bon_update}
        )
        
        updated_bon = await db.bons_travail.find_one({"id": bon_id})
        
        if "_id" in updated_bon:
            del updated_bon["_id"]
        
        return updated_bon
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur mise à jour bon {bon_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/bons-travail/{bon_id}", response_model=SuccessResponse)
async def delete_bon_travail(
    bon_id: str,
    current_user: dict = Depends(get_current_admin_user)
):
    """Supprimer un bon de travail"""
    try:
        bon = await db.bons_travail.find_one({"id": bon_id})
        if not bon:
            raise HTTPException(status_code=404, detail="Bon de travail non trouvé")
        
        await db.bons_travail.delete_one({"id": bon_id})
        
        return {"success": True, "message": "Bon de travail supprimé"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur suppression bon {bon_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== GÉNÉRATION PDF & EMAIL ====================

@router.get("/bons-travail/{bon_id}/pdf")
async def generate_bon_pdf(
    bon_id: str,
    token: str = None,
    current_user: dict = Depends(get_current_user_optional)
):
    """Générer un PDF (HTML) pour un bon de travail - Format MAINT_FE_004_V02"""
    try:
        # Vérifier l'authentification via token si nécessaire
        if not current_user and token:
            payload = decode_access_token(token)
            if payload is None:
                raise HTTPException(status_code=401, detail="Token invalide ou expiré")
        elif not current_user and not token:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        bon = await db.bons_travail.find_one({"id": bon_id})
        if not bon:
            raise HTTPException(status_code=404, detail="Bon de travail non trouvé")
        
        # Générer le HTML avec le template MAINT_FE_004_V02
        html_content = generate_bon_travail_html(bon)
        return HTMLResponse(content=html_content)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur génération PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/bons-travail/{bon_id}/email")
async def send_bon_email(
    bon_id: str,
    email_to: str,
    current_user: dict = Depends(get_current_user)
):
    """Envoyer un bon de travail par email"""
    try:
        bon = await db.bons_travail.find_one({"id": bon_id})
        if not bon:
            raise HTTPException(status_code=404, detail="Bon de travail non trouvé")
        
        # TODO: Implémenter l'envoi email avec SMTP
        # Pour l'instant, retourner un message
        
        return {
            "success": True,
            "message": "Envoi email en cours de développement",
            "bon_id": bon_id,
            "email_to": email_to
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur envoi email: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== FORM TEMPLATES ====================

@router.get("/form-templates")
async def get_form_templates(current_user: dict = Depends(get_current_user)):
    """Récupérer tous les modèles de formulaires (hors corbeille)"""
    try:
        # Exclure les templates en corbeille (soft-deleted)
        templates = await db.form_templates.find(
            {"$or": [{"deleted_at": None}, {"deleted_at": {"$exists": False}}]},
            {"_id": 0}
        ).to_list(length=None)
        
        # Templates système par défaut avec leurs champs
        BON_TRAVAIL_FIELDS = [
            {"label": "Localisation / Ligne", "type": "text", "required": True},
            {"label": "Description des travaux", "type": "textarea", "required": True},
            {"label": "Nom des intervenants", "type": "text", "required": True},
            {"label": "Type de travaux", "type": "select", "options": ["Mécanique", "Électrique", "Automatisme", "Plomberie", "Autre"]},
            {"label": "Risques identifiés", "type": "textarea"},
            {"label": "Consignations nécessaires", "type": "checkbox_group", "options": ["Consignation", "Déconsignation"]},
            {"label": "Précautions matérielles", "type": "checkbox_group", "options": ["Échafaudage", "Harnais", "Nacelle", "Ligne vie"]},
            {"label": "Équipements de Protection (EPI)", "type": "checkbox_group", "options": ["Casque", "Gants", "Lunettes", "Masque", "Chaussures S3", "Gilet HV", "Bouchons oreilles"]},
            {"label": "Précautions environnementales", "type": "checkbox_group", "options": ["Balisage", "Signalisation", "Permis feu", "Ventilation"]},
            {"label": "Date d'engagement", "type": "date", "required": True},
            {"label": "Nom agent de maîtrise", "type": "text", "required": True},
            {"label": "Nom du représentant", "type": "text", "required": True}
        ]
        
        AUTORISATION_FIELDS = [
            {"label": "Type d'autorisation", "type": "select", "options": ["Travail en hauteur", "Permis de feu", "Espace confiné", "Travail sur toiture", "Autre"], "required": True},
            {"label": "Lieu des travaux", "type": "text", "required": True},
            {"label": "Nature des travaux", "type": "textarea", "required": True},
            {"label": "Date de début", "type": "date", "required": True},
            {"label": "Date de fin", "type": "date", "required": True},
            {"label": "Entreprise intervenante", "type": "text"},
            {"label": "Nom du responsable", "type": "text", "required": True},
            {"label": "Mesures de prévention", "type": "textarea"},
            {"label": "Signature responsable", "type": "text", "required": True}
        ]
        
        # Si aucun template, créer les templates par défaut
        if not templates:
            default_templates = [
                {
                    "id": "default-bon-travail",
                    "nom": "Bon de travail",
                    "type": "BON_TRAVAIL",
                    "description": "Formulaire standard pour les bons de travail de maintenance",
                    "fields": BON_TRAVAIL_FIELDS,
                    "actif": True,
                    "is_system": True,
                    "created_at": datetime.now(timezone.utc).isoformat()
                },
                {
                    "id": "default-autorisation",
                    "nom": "Autorisation particulière",
                    "type": "AUTORISATION",
                    "description": "Formulaire standard pour les autorisations de travail spéciales",
                    "fields": AUTORISATION_FIELDS,
                    "actif": True,
                    "is_system": True,
                    "created_at": datetime.now(timezone.utc).isoformat()
                }
            ]
            for tpl in default_templates:
                await db.form_templates.insert_one(tpl)
            return default_templates
        
        # S'assurer que les templates système ont leurs champs
        for tpl in templates:
            if tpl.get("is_system") and not tpl.get("fields"):
                if tpl.get("type") == "BON_TRAVAIL":
                    tpl["fields"] = BON_TRAVAIL_FIELDS
                    await db.form_templates.update_one({"id": tpl["id"]}, {"$set": {"fields": BON_TRAVAIL_FIELDS}})
                elif tpl.get("type") == "AUTORISATION":
                    tpl["fields"] = AUTORISATION_FIELDS
                    await db.form_templates.update_one({"id": tpl["id"]}, {"$set": {"fields": AUTORISATION_FIELDS}})
        
        return templates
    except Exception as e:
        logger.error(f"Erreur récupération templates: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/form-templates/{template_id}")
async def get_form_template(template_id: str, current_user: dict = Depends(get_current_user)):
    """Récupérer un modèle de formulaire par ID"""
    try:
        template = await db.form_templates.find_one({"id": template_id}, {"_id": 0})
        if not template:
            raise HTTPException(status_code=404, detail="Modèle non trouvé")
        return template
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération template: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/form-templates")
async def create_form_template(
    template_data: dict,
    current_user: dict = Depends(get_current_admin_user)
):
    """Créer un nouveau modèle de formulaire (admin uniquement)"""
    try:
        template = {
            "id": str(uuid.uuid4()),
            "nom": template_data.get("nom"),
            "type": template_data.get("type", "CUSTOM"),
            "description": template_data.get("description", ""),
            "fields": template_data.get("fields", []),
            "actif": template_data.get("actif", True),
            "is_system": False,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "created_by": current_user.get("id")
        }
        
        await db.form_templates.insert_one(template)
        template.pop("_id", None)
        
        return template
    except Exception as e:
        logger.error(f"Erreur création template: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/form-templates/{template_id}")
async def update_form_template(
    template_id: str,
    template_data: dict,
    current_user: dict = Depends(get_current_admin_user)
):
    """Mettre à jour un modèle de formulaire (admin uniquement)"""
    try:
        existing = await db.form_templates.find_one({"id": template_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Modèle non trouvé")
        
        update_data = {
            "nom": template_data.get("nom", existing.get("nom")),
            "description": template_data.get("description", existing.get("description")),
            "fields": template_data.get("fields", existing.get("fields", [])),
            "actif": template_data.get("actif", existing.get("actif")),
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "updated_by": current_user.get("id")
        }
        # Ne pas changer le type pour les templates système
        if not existing.get("is_system"):
            update_data["type"] = template_data.get("type", existing.get("type"))
        
        await db.form_templates.update_one({"id": template_id}, {"$set": update_data})
        
        updated = await db.form_templates.find_one({"id": template_id}, {"_id": 0})
        return updated
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur mise à jour template: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/form-templates/{template_id}", response_model=SuccessResponse)
async def delete_form_template(
    template_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Supprimer un modèle de formulaire vers la corbeille (admin ou permission suppression documentations)"""
    try:
        # Vérification des droits : admin ou permission delete sur documentations
        is_admin = current_user.get("role") == "ADMIN"
        perms = current_user.get("permissions", {})
        can_delete = is_admin or perms.get("documentations", {}).get("delete", False)
        if not can_delete:
            raise HTTPException(status_code=403, detail="Permission refusée")

        existing = await db.form_templates.find_one({"id": template_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Modèle non trouvé")

        if existing.get("is_system"):
            raise HTTPException(status_code=400, detail="Les modèles système ne peuvent pas être supprimés")

        # Soft-delete : déplacer vers la corbeille
        now = datetime.now(timezone.utc)
        await db.form_templates.update_one(
            {"id": template_id},
            {"$set": {
                "deleted_at": now,
                "deleted_by": current_user.get("id"),
            }}
        )

        return {"success": True, "message": "Modèle déplacé dans la corbeille"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur suppression template: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== GÉNÉRATION IA DE FORMULAIRES ====================

FORM_AI_PROMPT = """Tu es un expert en création de formulaires pour une application GMAO (Gestion de Maintenance Assistée par Ordinateur) appelée FSAO Iris.

L'utilisateur te fournit une description, une image de formulaire papier, un fichier Excel, ou un prompt JSON. Tu dois analyser l'entrée et générer la structure JSON d'un modèle de formulaire.

=== TYPES DE CHAMPS DISPONIBLES ===
- "text" : Champ texte simple (une ligne)
- "textarea" : Zone de texte multiligne
- "number" : Champ numérique
- "date" : Sélecteur de date
- "select" : Liste déroulante (nécessite un tableau "options")
- "checkbox" : Case à cocher unique (oui/non)
- "checkbox_group" : Groupe de cases à cocher (nécessite un tableau "options")
- "radio" : Boutons radio (nécessite un tableau "options")
- "signature" : Zone de signature
- "file" : Upload de fichier
- "image" : Upload d'image

=== FORMAT DE SORTIE ===
Retourne UNIQUEMENT un JSON valide (pas de markdown, pas de texte avant/après) avec cette structure exacte :
{
  "nom": "Nom du formulaire",
  "description": "Description courte du formulaire",
  "fields": [
    {
      "label": "Nom du champ",
      "type": "text|textarea|number|date|select|checkbox|checkbox_group|radio|signature|file|image",
      "required": true/false,
      "options": ["option1", "option2"]  // uniquement pour select, checkbox_group, radio
    }
  ]
}

=== RÈGLES ===
- Analyse minutieusement l'entrée (image, texte, Excel)
- Identifie chaque champ/zone de saisie et son type le plus approprié
- Marque comme "required" les champs qui semblent obligatoires
- Pour les listes de choix, utilise "select" ou "checkbox_group" selon le contexte
- Pour les zones de texte longues (commentaires, descriptions), utilise "textarea"
- Pour les signatures, utilise "signature"
- Génère un nom et une description pertinents pour le formulaire
- Le JSON doit être valide et parsable directement
"""


@router.post("/form-templates/generate-ai")
async def generate_form_template_ai(
    description: Optional[str] = Form(None),
    json_prompt: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    current_user: dict = Depends(get_current_admin_user)
):
    """Générer un modèle de formulaire via IA à partir d'une description, image, Excel ou JSON"""
    try:
        from emergentintegrations.llm.chat import LlmChat, UserMessage, FileContentWithMimeType
        import tempfile

        # Récupérer la clé LLM
        api_key = os.environ.get("EMERGENT_LLM_KEY")
        if not api_key:
            gk = await db.global_settings.find_one({"key": "EMERGENT_LLM_KEY"})
            if gk and gk.get("value"):
                api_key = gk["value"]
        if not api_key:
            raise HTTPException(status_code=500, detail="Clé LLM non configurée")

        # Récupérer le modèle IA configuré
        ai_settings = await db.global_settings.find_one({"key": "form_ai_model"})
        provider = "openai"
        model = "gpt-4o"
        if ai_settings and ai_settings.get("value"):
            provider = ai_settings["value"].get("provider", "openai")
            model = ai_settings["value"].get("model", "gpt-4o")

        # Créer le chat IA
        chat = LlmChat(
            api_key=api_key,
            session_id=f"form-gen-{uuid.uuid4()}",
            system_message=FORM_AI_PROMPT
        ).with_model(provider, model)

        # Construire le message
        user_text = ""
        file_contents = []

        if json_prompt:
            user_text = f"Voici un prompt JSON à convertir en formulaire :\n```json\n{json_prompt}\n```\nGénère le formulaire correspondant."
        elif description:
            user_text = f"Crée un formulaire basé sur cette description :\n\n{description}"
        else:
            user_text = "Analyse le fichier ci-joint et crée un formulaire basé sur son contenu."

        # Traiter le fichier uploadé
        if file:
            content = await file.read()
            mime_type, _ = mimetypes.guess_type(file.filename) if file.filename else (None, None)
            mime_type = mime_type or file.content_type or "application/octet-stream"

            if mime_type.startswith("image/"):
                # Image : envoyer directement à l'IA
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(file.filename).suffix) as tmp:
                    tmp.write(content)
                    tmp_path = tmp.name
                file_contents.append(FileContentWithMimeType(file_path=tmp_path, mime_type=mime_type))
                user_text += f"\n\nAnalyse cette image de formulaire et extrais tous les champs visibles."

            elif "spreadsheet" in mime_type or "excel" in mime_type or file.filename.endswith(('.xlsx', '.xls', '.csv')):
                # Excel/CSV : extraire le contenu textuel
                import io
                try:
                    import openpyxl
                    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
                    ws = wb.active
                    rows_text = []
                    for row in ws.iter_rows(max_row=50, values_only=True):
                        row_vals = [str(c) if c is not None else "" for c in row]
                        if any(v.strip() for v in row_vals):
                            rows_text.append(" | ".join(row_vals))
                    excel_text = "\n".join(rows_text)
                    user_text += f"\n\nContenu du fichier Excel :\n```\n{excel_text}\n```\nAnalyse ce contenu et crée un formulaire avec les champs appropriés."
                except Exception:
                    try:
                        text_content = content.decode('utf-8', errors='replace')
                        user_text += f"\n\nContenu du fichier :\n```\n{text_content[:5000]}\n```\nAnalyse et crée un formulaire."
                    except Exception:
                        user_text += "\n\nImpossible de lire le fichier. Crée un formulaire générique de maintenance."

            elif "pdf" in mime_type or file.filename.endswith('.pdf'):
                # PDF : extraire le texte
                try:
                    import io
                    from PyPDF2 import PdfReader
                    reader = PdfReader(io.BytesIO(content))
                    pages_text = []
                    for page in reader.pages[:20]:
                        text = page.extract_text()
                        if text and text.strip():
                            pages_text.append(text.strip())
                    pdf_text = "\n---\n".join(pages_text)
                    if pdf_text.strip():
                        user_text += f"\n\nContenu du fichier PDF :\n```\n{pdf_text[:8000]}\n```\nAnalyse ce contenu et crée un formulaire avec les champs appropriés."
                    else:
                        user_text += "\n\nLe PDF ne contient pas de texte extractible. Crée un formulaire basé sur la description fournie ou sur le nom du fichier."
                except Exception as pdf_err:
                    logger.warning(f"Erreur extraction PDF: {pdf_err}")
                    user_text += "\n\nImpossible de lire le PDF. Crée un formulaire basé sur la description fournie."

            elif "word" in mime_type or "document" in mime_type or file.filename.endswith(('.docx', '.doc')):
                # Word : extraire le texte
                try:
                    import io
                    from docx import Document as DocxDocument
                    doc = DocxDocument(io.BytesIO(content))
                    paragraphs = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            paragraphs.append(para.text.strip())
                    # Extraire aussi les tableaux
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                paragraphs.append(row_text)
                    word_text = "\n".join(paragraphs)
                    user_text += f"\n\nContenu du fichier Word :\n```\n{word_text[:8000]}\n```\nAnalyse ce contenu et crée un formulaire avec les champs appropriés."
                except Exception as docx_err:
                    logger.warning(f"Erreur extraction Word: {docx_err}")
                    user_text += "\n\nImpossible de lire le fichier Word. Crée un formulaire basé sur la description fournie."

            else:
                # Autre type de fichier
                try:
                    text_content = content.decode('utf-8', errors='replace')
                    user_text += f"\n\nContenu du fichier :\n```\n{text_content[:5000]}\n```"
                except Exception:
                    user_text += "\n\nFichier non lisible."

        if not user_text.strip() and not file:
            raise HTTPException(status_code=400, detail="Fournissez une description, un JSON ou un fichier")

        # Appeler l'IA
        msg = UserMessage(text=user_text, file_contents=file_contents if file_contents else None)
        response = await chat.send_message(msg)

        # Parser la réponse JSON
        import re
        json_match = re.search(r'\{[\s\S]*\}', response)
        if not json_match:
            raise HTTPException(status_code=500, detail="L'IA n'a pas retourné un JSON valide")

        result = json.loads(json_match.group())

        # Valider la structure
        if "fields" not in result or not isinstance(result["fields"], list):
            raise HTTPException(status_code=500, detail="Structure de formulaire invalide")

        return {
            "success": True,
            "template": {
                "nom": result.get("nom", "Formulaire généré par IA"),
                "description": result.get("description", ""),
                "fields": result["fields"]
            }
        }
    except HTTPException:
        raise
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="L'IA n'a pas retourné un JSON valide. Réessayez.")
    except Exception as e:
        logger.error(f"Erreur génération IA formulaire: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== PARAMÈTRES MODÈLE IA FORMULAIRES ====================

@router.get("/ai-model-config")
async def get_ai_model_config(current_user: dict = Depends(get_current_user)):
    """Récupérer la configuration du modèle IA pour les formulaires"""
    try:
        config = await db.global_settings.find_one({"key": "form_ai_model"}, {"_id": 0})
        if config and config.get("value"):
            return config["value"]
        return {"provider": "openai", "model": "gpt-4o"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/ai-model-config")
async def update_ai_model_config(
    data: dict,
    current_user: dict = Depends(get_current_admin_user)
):
    """Mettre à jour la configuration du modèle IA pour les formulaires"""
    try:
        provider = data.get("provider", "openai")
        model = data.get("model", "gpt-4o")
        
        await db.global_settings.update_one(
            {"key": "form_ai_model"},
            {"$set": {"key": "form_ai_model", "value": {"provider": provider, "model": model}, "updated_at": datetime.now(timezone.utc).isoformat()}},
            upsert=True
        )
        return {"success": True, "provider": provider, "model": model}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ==================== CUSTOM FORM INSTANCES (Filled Forms) ====================

@router.get("/custom-forms")
async def get_custom_forms(
    pole_id: Optional[str] = None,
    template_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer les formulaires personnalisés remplis"""
    try:
        query = {}
        if pole_id:
            query["pole_id"] = pole_id
        if template_id:
            query["template_id"] = template_id
        
        forms = await db.custom_forms.find(query, {"_id": 0}).to_list(length=None)
        return forms
    except Exception as e:
        logger.error(f"Erreur récupération custom forms: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/custom-forms/{form_id}")
async def get_custom_form(form_id: str, current_user: dict = Depends(get_current_user)):
    """Récupérer un formulaire personnalisé par ID"""
    try:
        form = await db.custom_forms.find_one({"id": form_id}, {"_id": 0})
        if not form:
            raise HTTPException(status_code=404, detail="Formulaire non trouvé")
        return form
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération custom form: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/custom-forms")
async def create_custom_form(
    form_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Créer un nouveau formulaire personnalisé rempli"""
    try:
        # Vérifier que le template existe
        template = await db.form_templates.find_one({"id": form_data.get("template_id")})
        if not template:
            raise HTTPException(status_code=404, detail="Modèle de formulaire non trouvé")
        
        custom_form = {
            "id": str(uuid.uuid4()),
            "template_id": form_data.get("template_id"),
            "template_name": template.get("nom"),
            "pole_id": form_data.get("pole_id"),
            "titre": form_data.get("titre", template.get("nom")),
            "field_values": form_data.get("field_values", {}),
            "attachments": form_data.get("attachments", []),
            "signature_data": form_data.get("signature_data"),
            "logo_url": form_data.get("logo_url"),
            "status": form_data.get("status", "BROUILLON"),
            "created_at": datetime.now(timezone.utc).isoformat(),
            "created_by": current_user.get("id"),
            "created_by_name": f"{current_user.get('prenom', '')} {current_user.get('nom', '')}".strip()
        }
        
        await db.custom_forms.insert_one(custom_form)
        custom_form.pop("_id", None)
        
        return custom_form
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur création custom form: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/custom-forms/{form_id}")
async def update_custom_form(
    form_id: str,
    form_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Mettre à jour un formulaire personnalisé"""
    try:
        existing = await db.custom_forms.find_one({"id": form_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Formulaire non trouvé")
        
        # Vérifier les permissions
        is_admin = current_user.get("role") == "ADMIN"
        is_creator = existing.get("created_by") == current_user.get("id")
        
        if not is_admin and not is_creator:
            raise HTTPException(status_code=403, detail="Non autorisé à modifier ce formulaire")
        
        update_data = {
            "titre": form_data.get("titre", existing.get("titre")),
            "field_values": form_data.get("field_values", existing.get("field_values")),
            "attachments": form_data.get("attachments", existing.get("attachments")),
            "signature_data": form_data.get("signature_data", existing.get("signature_data")),
            "logo_url": form_data.get("logo_url", existing.get("logo_url")),
            "status": form_data.get("status", existing.get("status")),
            "updated_at": datetime.now(timezone.utc).isoformat(),
            "updated_by": current_user.get("id")
        }
        
        await db.custom_forms.update_one({"id": form_id}, {"$set": update_data})
        
        updated = await db.custom_forms.find_one({"id": form_id}, {"_id": 0})
        return updated
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur mise à jour custom form: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/custom-forms/{form_id}", response_model=SuccessResponse)
async def delete_custom_form(
    form_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Supprimer un formulaire personnalisé"""
    try:
        existing = await db.custom_forms.find_one({"id": form_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Formulaire non trouvé")
        
        # Vérifier les permissions
        is_admin = current_user.get("role") == "ADMIN"
        is_creator = existing.get("created_by") == current_user.get("id")
        
        if not is_admin and not is_creator:
            raise HTTPException(status_code=403, detail="Non autorisé à supprimer ce formulaire")
        
        await db.custom_forms.delete_one({"id": form_id})
        
        return {"success": True, "message": "Formulaire supprimé"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur suppression custom form: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/bons-de-travail/save")
async def save_bon_de_travail_v2(
    payload: dict,
    current_user: dict = Depends(get_current_user)
):
    """
    Sauvegarde un Bon de Travail au format MAINT/FE/004 V2 (champs booléens individuels).
    Compatible avec l'explorateur documentations (pole_id obligatoire).
    Retourne l'id du bon créé ou mis à jour.
    """
    import uuid as _uuid
    try:
        pole_id = payload.get("pole_id")
        if not pole_id:
            raise HTTPException(status_code=400, detail="pole_id est requis pour enregistrer un bon de travail")

        bon_id = payload.get("id") or str(_uuid.uuid4())
        now = datetime.now(timezone.utc).isoformat()

        doc = {
            "id": bon_id,
            "pole_id": pole_id,
            "form_version": 2,
            # ── Champs affichés dans l'explorateur ──────────────────
            "localisation_ligne": payload.get("localisation", ""),
            "description_travaux": payload.get("description", ""),
            "nom_intervenants": payload.get("intervenants", ""),
            "titre": payload.get("titre") or payload.get("localisation") or "Bon de travail",
            "statut": payload.get("statut", "BROUILLON"),
            "entreprise": payload.get("entreprise", "Non assignée"),
            # ── Champs booléens complets (pour régénérer le PDF) ────
            "form_data": {k: v for k, v in payload.items() if k not in ("pole_id", "id", "titre")},
            "created_at": payload.get("created_at", now),
            "updated_at": now,
            "created_by": current_user.get("id"),
        }

        existing = await db.bons_travail.find_one({"id": bon_id})
        if existing:
            await db.bons_travail.update_one({"id": bon_id}, {"$set": doc})
        else:
            await db.bons_travail.insert_one(doc)

        doc.pop("_id", None)
        return {"id": bon_id, "status": "ok", "titre": doc["titre"]}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur sauvegarde bon de travail v2: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/bons-de-travail/generate-pdf")
async def generate_bon_de_travail_pdf_endpoint(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """
    Génère un PDF ReportLab du Bon de Travail MAINT/FE/004 V2.
    Accepte un dict avec tous les champs du formulaire.
    Fonctionne pour un bon vierge (dict vide) ou pré-rempli.
    """
    try:
        from bon_de_travail_reportlab import generate_bon_travail_pdf
        pdf_bytes = generate_bon_travail_pdf(data)
        from fastapi.responses import Response
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": "inline; filename=bon_de_travail_MAINT_FE_004.pdf"}
        )
    except Exception as e:
        logger.error(f"Erreur génération PDF bon de travail: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/custom-forms/{form_id}/pdf")
async def generate_custom_form_pdf(
    form_id: str,
    token: Optional[str] = None,
    current_user: dict = Depends(get_current_user_optional)
):
    """Générer un PDF pour un formulaire personnalisé"""
    try:
        # Vérifier l'authentification
        if not current_user and token:
            payload = decode_access_token(token)
            if payload is None:
                raise HTTPException(status_code=401, detail="Token invalide")
        elif not current_user and not token:
            raise HTTPException(status_code=401, detail="Non authentifié")
        
        # Récupérer le formulaire
        form = await db.custom_forms.find_one({"id": form_id}, {"_id": 0})
        if not form:
            raise HTTPException(status_code=404, detail="Formulaire non trouvé")
        
        # Récupérer le template
        template = await db.form_templates.find_one({"id": form.get("template_id")}, {"_id": 0})
        
        # Générer le HTML
        html_content = generate_custom_form_html(form, template)
        
        return HTMLResponse(content=html_content)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur génération PDF custom form: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def generate_custom_form_html(form: dict, template: dict) -> str:
    """Génère le HTML pour un formulaire personnalisé"""
    fields = template.get("fields", []) if template else []
    field_values = form.get("field_values", {})
    
    # Générer les lignes de champs
    fields_html = ""
    for field in fields:
        field_id = field.get("id")
        field_label = field.get("label", "")
        field_type = field.get("type", "text")
        value = field_values.get(field_id, "")
        
        # Formatage selon le type
        if field_type == "checkbox":
            value = "✓ Oui" if value else "✗ Non"
        elif field_type == "switch":
            value = "✓ Oui" if value else "✗ Non"
        elif field_type == "select":
            # La valeur est déjà le label sélectionné
            pass
        elif field_type == "date" and value:
            try:
                from datetime import datetime as dt
                value = dt.fromisoformat(value.replace('Z', '+00:00')).strftime("%d/%m/%Y")
            except:
                pass
        elif field_type == "textarea":
            value = value.replace('\n', '<br>') if value else ""
        
        fields_html += f"""
        <tr>
            <td style="padding: 8px; border: 1px solid #ddd; background: #f9f9f9; font-weight: 500; width: 30%;">{field_label}</td>
            <td style="padding: 8px; border: 1px solid #ddd;">{value or '-'}</td>
        </tr>
        """
    
    # Signature
    signature_html = ""
    if form.get("signature_data"):
        signature_html = f"""
        <div style="margin-top: 30px; page-break-inside: avoid;">
            <h3 style="color: #333; border-bottom: 2px solid #2563eb; padding-bottom: 5px;">Signature</h3>
            <img src="{form.get('signature_data')}" style="max-width: 300px; border: 1px solid #ddd; padding: 10px;" />
        </div>
        """
    
    # Logo
    logo_html = ""
    if form.get("logo_url"):
        logo_html = f'<img src="{form.get("logo_url")}" style="max-height: 60px;" />'
    
    html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{form.get('titre', 'Formulaire')}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; color: #333; }}
            .header {{ display: flex; justify-content: space-between; align-items: center; margin-bottom: 30px; border-bottom: 3px solid #2563eb; padding-bottom: 15px; }}
            .title {{ font-size: 24px; font-weight: bold; color: #2563eb; }}
            .meta {{ color: #666; font-size: 12px; margin-top: 5px; }}
            table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
            @media print {{
                body {{ margin: 10mm; }}
                .no-print {{ display: none; }}
            }}
        </style>
    </head>
    <body>
        <div class="header">
            <div>
                <div class="title">{form.get('titre', 'Formulaire personnalisé')}</div>
                <div class="meta">
                    Créé le {form.get('created_at', '')[:10]} par {form.get('created_by_name', 'Inconnu')}
                </div>
            </div>
            {logo_html}
        </div>
        
        <table>
            {fields_html}
        </table>
        
        {signature_html}
        
        <div style="margin-top: 50px; text-align: center; color: #999; font-size: 10px;">
            Document généré le {datetime.now(timezone.utc).strftime('%d/%m/%Y à %H:%M')}
        </div>
    </body>
    </html>
    """
    
    return html


# ==================== DOSSIERS (Vue Explorateur) ====================

@router.get("/poles/{pole_id}/folders")
async def get_folders(
    pole_id: str,
    parent_id: Optional[str] = None,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer les dossiers d'un pôle (optionnellement dans un parent)"""
    try:
        query = {"pole_id": pole_id, "parent_id": parent_id}
        folders = await db.doc_folders.find(query, {"_id": 0}).to_list(length=None)
        return folders
    except Exception as e:
        logger.error(f"Erreur récupération dossiers: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/poles/{pole_id}/folders")
async def create_folder(
    pole_id: str,
    folder_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Créer un nouveau dossier dans un pôle"""
    try:
        pole = await db.poles_service.find_one({"id": pole_id})
        if not pole:
            raise HTTPException(status_code=404, detail="Pôle non trouvé")

        folder = {
            "id": str(uuid.uuid4()),
            "pole_id": pole_id,
            "parent_id": folder_data.get("parent_id"),
            "name": folder_data.get("name", "Nouveau dossier"),
            "created_at": datetime.now(timezone.utc).isoformat(),
            "created_by": current_user.get("id"),
            "updated_at": datetime.now(timezone.utc).isoformat()
        }

        await db.doc_folders.insert_one(folder)
        folder.pop("_id", None)

        if realtime_manager:
            await realtime_manager.emit_event("documentations", "created", folder, user_id=current_user["id"])

        return folder
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur création dossier: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/folders/{folder_id}")
async def update_folder(
    folder_id: str,
    folder_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Renommer ou déplacer un dossier"""
    try:
        existing = await db.doc_folders.find_one({"id": folder_id})
        if not existing:
            raise HTTPException(status_code=404, detail="Dossier non trouvé")

        update_data = {}
        if "name" in folder_data:
            update_data["name"] = folder_data["name"]
        if "parent_id" in folder_data:
            update_data["parent_id"] = folder_data["parent_id"]
        update_data["updated_at"] = datetime.now(timezone.utc).isoformat()

        await db.doc_folders.update_one({"id": folder_id}, {"$set": update_data})
        updated = await db.doc_folders.find_one({"id": folder_id}, {"_id": 0})

        if realtime_manager:
            await realtime_manager.emit_event("documentations", "updated", updated, user_id=current_user["id"])

        return updated
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur mise à jour dossier: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.delete("/folders/{folder_id}", response_model=SuccessResponse)
async def delete_folder(
    folder_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Supprimer un dossier et déplacer son contenu au parent"""
    try:
        folder = await db.doc_folders.find_one({"id": folder_id})
        if not folder:
            raise HTTPException(status_code=404, detail="Dossier non trouvé")

        parent_id = folder.get("parent_id")

        # Déplacer les sous-dossiers vers le parent
        await db.doc_folders.update_many(
            {"parent_id": folder_id},
            {"$set": {"parent_id": parent_id}}
        )
        # Déplacer les documents vers le parent
        await db.documents.update_many(
            {"folder_id": folder_id},
            {"$set": {"folder_id": parent_id}}
        )

        await db.doc_folders.delete_one({"id": folder_id})

        if realtime_manager:
            await realtime_manager.emit_event("documentations", "deleted", {"id": folder_id}, user_id=current_user["id"])

        return {"success": True, "message": "Dossier supprimé"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur suppression dossier: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.put("/documents/{document_id}/move")
async def move_document(
    document_id: str,
    move_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Déplacer un document vers un dossier ou un autre pôle"""
    try:
        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")

        update = {"updated_at": datetime.now(timezone.utc).isoformat()}
        if "folder_id" in move_data:
            update["folder_id"] = move_data["folder_id"]
        if "pole_id" in move_data:
            update["pole_id"] = move_data["pole_id"]

        await db.documents.update_one({"id": document_id}, {"$set": update})
        updated = await db.documents.find_one({"id": document_id}, {"_id": 0})

        if realtime_manager:
            await realtime_manager.emit_event("documentations", "updated", updated, user_id=current_user["id"])

        return updated
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur déplacement document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/poles/{pole_id}/explorer")
async def get_explorer_contents(
    pole_id: str,
    folder_id: Optional[str] = None,
    sort_by: Optional[str] = "name",
    current_user: dict = Depends(get_current_user)
):
    """Récupérer le contenu d'un dossier pour la vue explorateur avec filtrage par permissions"""
    try:
        pole = await db.poles_service.find_one({"id": pole_id})
        if not pole:
            raise HTTPException(status_code=404, detail="Pôle non trouvé")
        if "_id" in pole:
            del pole["_id"]

        user_role = current_user.get("role", "")
        user_service = current_user.get("service", "")
        is_admin = user_role == "ADMIN"
        is_maintenance = user_service == "Maintenance" or pole.get("pole") == "MAINTENANCE"

        # Sous-dossiers
        folders = await db.doc_folders.find(
            {"pole_id": pole_id, "parent_id": folder_id}, {"_id": 0}
        ).to_list(length=None)

        # Filtrer par permissions
        if not is_admin:
            folders = [f for f in folders if not f.get("hidden_for_users", False)]
        if not is_admin and not is_maintenance:
            folders = [f for f in folders if not f.get("hidden_for_external", False)]

        # Documents dans ce dossier
        if folder_id is None:
            documents = await db.documents.find(
                {"pole_id": pole_id, "$or": [{"folder_id": None}, {"folder_id": {"$exists": False}}]}, {"_id": 0}
            ).to_list(length=None)
        else:
            documents = await db.documents.find(
                {"pole_id": pole_id, "folder_id": folder_id}, {"_id": 0}
            ).to_list(length=None)

        # Filtrer par permissions
        if not is_admin:
            documents = [d for d in documents if not d.get("hidden_for_users", False)]
        if not is_admin and not is_maintenance:
            documents = [d for d in documents if not d.get("hidden_for_external", False)]

        # Bons de travail (uniquement à la racine du pôle)
        bons_travail = []
        if folder_id is None:
            bons_travail = await db.bons_travail.find(
                {"pole_id": pole_id}, {"_id": 0}
            ).to_list(length=None)

        # Tri
        sort_key_map = {
            "name": lambda x: (x.get("name") or x.get("titre") or x.get("fichier_nom") or "").lower(),
            "date": lambda x: x.get("created_at") or x.get("updated_at") or "",
            "type": lambda x: x.get("fichier_type") or x.get("type_document") or ""
        }
        sort_fn = sort_key_map.get(sort_by, sort_key_map["name"])
        folders.sort(key=sort_fn)
        documents.sort(key=sort_fn)
        bons_travail.sort(key=sort_fn)

        # Breadcrumb
        breadcrumb = [{"id": pole_id, "name": pole.get("nom", ""), "type": "pole"}]
        if folder_id:
            current_folder_id = folder_id
            while current_folder_id:
                f = await db.doc_folders.find_one({"id": current_folder_id}, {"_id": 0})
                if f:
                    breadcrumb.insert(1, {"id": f["id"], "name": f["name"], "type": "folder"})
                    current_folder_id = f.get("parent_id")
                else:
                    break

        return {
            "pole": pole,
            "folders": folders,
            "documents": documents,
            "bons_travail": bons_travail,
            "breadcrumb": breadcrumb,
            "current_folder_id": folder_id
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération contenu explorateur: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== COPIER / COUPER-COLLER ====================

@router.post("/copy")
async def copy_node(
    copy_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Copier un document ou dossier vers un emplacement cible"""
    try:
        node_id = copy_data.get("node_id")
        node_type = copy_data.get("node_type")  # "document" ou "folder"
        target_pole_id = copy_data.get("target_pole_id")
        target_folder_id = copy_data.get("target_folder_id")

        if node_type == "document":
            doc = await db.documents.find_one({"id": node_id}, {"_id": 0})
            if not doc:
                raise HTTPException(status_code=404, detail="Document non trouvé")
            # Créer une copie
            new_doc = {**doc}
            new_doc["id"] = str(uuid.uuid4())
            new_doc["titre"] = f"{doc.get('titre', 'Document')} (copie)"
            if doc.get("fichier_nom"):
                name, ext = os.path.splitext(doc["fichier_nom"])
                new_doc["fichier_nom"] = f"{name} (copie){ext}"
            if target_pole_id:
                new_doc["pole_id"] = target_pole_id
            new_doc["folder_id"] = target_folder_id
            new_doc["created_at"] = datetime.now(timezone.utc).isoformat()
            new_doc["created_by"] = current_user.get("id")
            # Copier le fichier physique si existant
            if doc.get("fichier_url"):
                src_path = Path(f"/app/backend{doc['fichier_url']}")
                if src_path.exists():
                    upload_dir = Path("uploads/documents")
                    upload_dir.mkdir(parents=True, exist_ok=True)
                    new_filename = f"{new_doc['id']}_{uuid.uuid4()}{src_path.suffix}"
                    dst_path = upload_dir / new_filename
                    import shutil
                    shutil.copy2(src_path, dst_path)
                    new_doc["fichier_url"] = f"/uploads/documents/{new_filename}"
            await db.documents.insert_one(new_doc)
            new_doc.pop("_id", None)
            if realtime_manager:
                await realtime_manager.emit_event("documentations", "created", new_doc, user_id=current_user["id"])
            # Audit
            await audit_service.log_action(
                user_id=current_user["id"],
                user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
                user_email=current_user.get("email", ""),
                action=ActionType.COPY,
                entity_type=EntityType.DOCUMENTATION,
                entity_id=new_doc["id"],
                entity_name=doc.get("fichier_nom") or doc.get("titre", "Document"),
                details=f"a copié le document \"{doc.get('fichier_nom') or doc.get('titre', 'Document')}\""
            )
            return new_doc

        elif node_type == "folder":
            folder = await db.doc_folders.find_one({"id": node_id}, {"_id": 0})
            if not folder:
                raise HTTPException(status_code=404, detail="Dossier non trouvé")
            new_folder = {**folder}
            new_folder["id"] = str(uuid.uuid4())
            new_folder["name"] = f"{folder.get('name', 'Dossier')} (copie)"
            if target_pole_id:
                new_folder["pole_id"] = target_pole_id
            new_folder["parent_id"] = target_folder_id
            new_folder["created_at"] = datetime.now(timezone.utc).isoformat()
            new_folder["created_by"] = current_user.get("id")
            await db.doc_folders.insert_one(new_folder)
            new_folder.pop("_id", None)
            if realtime_manager:
                await realtime_manager.emit_event("documentations", "created", new_folder, user_id=current_user["id"])
            # Audit
            await audit_service.log_action(
                user_id=current_user["id"],
                user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
                user_email=current_user.get("email", ""),
                action=ActionType.COPY,
                entity_type=EntityType.DOCUMENTATION,
                entity_id=new_folder["id"],
                entity_name=folder.get("name", "Dossier"),
                details=f"a copié le dossier \"{folder.get('name', 'Dossier')}\""
            )
            return new_folder
        else:
            raise HTTPException(status_code=400, detail="Type de noeud invalide")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur copie: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/move")
async def move_node(
    move_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Déplacer (couper-coller) un document ou dossier"""
    try:
        node_id = move_data.get("node_id")
        node_type = move_data.get("node_type")
        target_pole_id = move_data.get("target_pole_id")
        target_folder_id = move_data.get("target_folder_id")

        if node_type == "document":
            doc = await db.documents.find_one({"id": node_id})
            if not doc:
                raise HTTPException(status_code=404, detail="Document non trouvé")
            update = {"updated_at": datetime.now(timezone.utc).isoformat(), "folder_id": target_folder_id}
            if target_pole_id:
                update["pole_id"] = target_pole_id
            await db.documents.update_one({"id": node_id}, {"$set": update})
            updated = await db.documents.find_one({"id": node_id}, {"_id": 0})
            if realtime_manager:
                await realtime_manager.emit_event("documentations", "updated", updated, user_id=current_user["id"])
            # Audit
            await audit_service.log_action(
                user_id=current_user["id"],
                user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
                user_email=current_user.get("email", ""),
                action=ActionType.MOVE,
                entity_type=EntityType.DOCUMENTATION,
                entity_id=node_id,
                entity_name=doc.get("fichier_nom") or doc.get("titre", "Document"),
                details=f"a déplacé le document \"{doc.get('fichier_nom') or doc.get('titre', 'Document')}\""
            )
            return updated

        elif node_type == "folder":
            folder = await db.doc_folders.find_one({"id": node_id})
            if not folder:
                raise HTTPException(status_code=404, detail="Dossier non trouvé")
            update = {"updated_at": datetime.now(timezone.utc).isoformat(), "parent_id": target_folder_id}
            if target_pole_id:
                update["pole_id"] = target_pole_id
            await db.doc_folders.update_one({"id": node_id}, {"$set": update})
            updated = await db.doc_folders.find_one({"id": node_id}, {"_id": 0})
            if realtime_manager:
                await realtime_manager.emit_event("documentations", "updated", updated, user_id=current_user["id"])
            # Audit
            await audit_service.log_action(
                user_id=current_user["id"],
                user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
                user_email=current_user.get("email", ""),
                action=ActionType.MOVE,
                entity_type=EntityType.DOCUMENTATION,
                entity_id=node_id,
                entity_name=folder.get("name", "Dossier"),
                details=f"a déplacé le dossier \"{folder.get('name', 'Dossier')}\""
            )
            return updated
        else:
            raise HTTPException(status_code=400, detail="Type de noeud invalide")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur déplacement: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== PERMISSIONS ====================

@router.patch("/permissions/{node_id}")
async def toggle_permissions(
    node_id: str,
    perm_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Basculer la visibilité d'un document ou dossier"""
    try:
        node_type = perm_data.get("node_type", "document")
        field = perm_data.get("field")  # "hidden_for_external" ou "hidden_for_users"

        if field not in ("hidden_for_external", "hidden_for_users"):
            raise HTTPException(status_code=400, detail="Champ de permission invalide")

        collection = db.documents if node_type == "document" else db.doc_folders
        node = await collection.find_one({"id": node_id})
        if not node:
            raise HTTPException(status_code=404, detail="Élément non trouvé")

        current_value = node.get(field, False)
        new_value = not current_value

        await collection.update_one(
            {"id": node_id},
            {"$set": {field: new_value, "updated_at": datetime.now(timezone.utc).isoformat()}}
        )

        updated = await collection.find_one({"id": node_id}, {"_id": 0})
        if realtime_manager:
            await realtime_manager.emit_event("documentations", "updated", updated, user_id=current_user["id"])

        # Audit
        node_name = node.get("fichier_nom") or node.get("titre") or node.get("name", "Élément")
        field_label = "services externes" if field == "hidden_for_external" else "utilisateurs"
        action_label = "masqué" if new_value else "rendu visible"
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
            user_email=current_user.get("email", ""),
            action=ActionType.PERMISSION_CHANGE,
            entity_type=EntityType.DOCUMENTATION,
            entity_id=node_id,
            entity_name=node_name,
            details=f"a {action_label} \"{node_name}\" aux {field_label}"
        )

        return {"success": True, "field": field, "value": new_value, "node": updated}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur permissions: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== ENVOYER VERS (autre pôle) ====================

@router.post("/send-to")
async def send_to_pole(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Copier un fichier/dossier vers la racine d'un autre pôle"""
    try:
        node_id = data.get("node_id")
        node_type = data.get("node_type")
        target_pole_id = data.get("target_pole_id")

        # Vérifier que le pôle cible existe
        target_pole = await db.poles_service.find_one({"id": target_pole_id})
        if not target_pole:
            raise HTTPException(status_code=404, detail="Pôle cible non trouvé")

        # Copier vers la racine du pôle cible
        result = await copy_node(
            {"node_id": node_id, "node_type": node_type, "target_pole_id": target_pole_id, "target_folder_id": None},
            current_user
        )
        # Audit (note: copy_node a déjà un audit de copie, ici on ajoute le contexte "envoyé vers")
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
            user_email=current_user.get("email", ""),
            action=ActionType.COPY,
            entity_type=EntityType.DOCUMENTATION,
            entity_id=node_id,
            details=f"a envoyé vers le pôle \"{target_pole.get('nom')}\""
        )
        return {"success": True, "message": f"Copié vers {target_pole.get('nom')}", "node": result}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur envoi vers pôle: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== PARTAGER PAR FSAO (EMAIL SMTP) ====================

@router.post("/share-email")
async def share_by_email(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Envoyer un document par email via le SMTP configuré dans l'application"""
    try:
        import email_service

        document_id = data.get("document_id")
        recipient = data.get("recipient")
        subject = data.get("subject", "Document partagé via FSAO")
        message = data.get("message", "")
        signature = data.get("signature", f"Cordialement,\n{current_user.get('prenom', '')} {current_user.get('nom', '')}\nFSAO Atlas")

        if not recipient:
            raise HTTPException(status_code=400, detail="Destinataire requis")

        doc = await db.documents.find_one({"id": document_id})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")

        # Construire le HTML de l'email
        html_content = f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
            <h2 style="color: #2563eb;">Document partagé</h2>
            <p>{message.replace(chr(10), '<br>')}</p>
            <div style="background: #f3f4f6; padding: 15px; border-radius: 8px; margin: 20px 0;">
                <p style="margin: 0;"><strong>Document :</strong> {doc.get('fichier_nom') or doc.get('titre', 'Document')}</p>
                {f"<p style='margin: 5px 0 0;'><strong>Type :</strong> {doc.get('fichier_type', 'N/A')}</p>" if doc.get('fichier_type') else ""}
            </div>
            <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 20px 0;" />
            <p style="white-space: pre-line; color: #6b7280;">{signature}</p>
        </div>
        """

        # Envoyer avec ou sans pièce jointe
        attachment_data = None
        attachment_filename = None
        if doc.get("fichier_url"):
            file_path = Path(f"/app/backend{doc['fichier_url']}")
            if file_path.exists():
                with open(file_path, "rb") as f:
                    attachment_data = f.read()
                attachment_filename = doc.get("fichier_nom", "document")

        if attachment_data:
            success = email_service.send_email_with_attachment(
                to_email=recipient,
                subject=subject,
                html_content=html_content,
                attachment_data=attachment_data,
                attachment_filename=attachment_filename
            )
        else:
            success = email_service.send_email(
                to_email=recipient,
                subject=subject,
                html_content=html_content
            )

        if success:
            # Audit
            doc_name = doc.get('fichier_nom') or doc.get('titre', 'Document')
            await audit_service.log_action(
                user_id=current_user["id"],
                user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
                user_email=current_user.get("email", ""),
                action=ActionType.SHARE,
                entity_type=EntityType.DOCUMENTATION,
                entity_id=document_id,
                entity_name=doc_name,
                details=f"a partagé \"{doc_name}\" par email FSAO à {recipient}"
            )
            return {"success": True, "message": f"Email envoyé à {recipient}"}
        else:
            raise HTTPException(status_code=500, detail="Échec de l'envoi de l'email")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur partage email: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


# ==================== INSÉRER DANS OT / AMÉLIORATION / M.PREV ====================

@router.get("/insert-targets")
async def get_insert_targets(
    target_type: str,
    current_user: dict = Depends(get_current_user)
):
    """Récupérer la liste des OT, Améliorations ou M.Prev pour le dialogue d'insertion"""
    try:
        if target_type == "work_order":
            items = await db.work_orders.find(
                {"statut": {"$in": ["OUVERT", "EN_COURS", "EN_ATTENTE"]}},
                {"_id": 0, "id": 1, "numero": 1, "titre": 1, "statut": 1}
            ).to_list(500)
            # S'assurer que chaque item a un champ 'id'
            for item in items:
                if not item.get("id"):
                    item["id"] = str(item.get("numero", ""))
            return items
        elif target_type == "improvement":
            items = await db.improvements.find(
                {"statut": {"$in": ["OUVERT", "EN_COURS", "EN_ATTENTE", "PROPOSEE"]}},
                {"_id": 0, "id": 1, "numero": 1, "titre": 1, "statut": 1}
            ).to_list(500)
            return items
        elif target_type == "preventive_maintenance":
            items = await db.preventive_maintenances.find(
                {"statut": "ACTIF"},
                {"_id": 0, "id": 1, "titre": 1, "statut": 1, "equipement_nom": 1}
            ).to_list(500)
            # Ajouter un label pour l'affichage
            for item in items:
                if not item.get("titre"):
                    item["titre"] = item.get("equipement_nom", "M.Prev")
            return items
        else:
            raise HTTPException(status_code=400, detail="Type cible invalide")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur récupération cibles insertion: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/insert-into")
async def insert_document_into(
    data: dict,
    current_user: dict = Depends(get_current_user)
):
    """Insérer un document dans un OT, une Amélioration ou une M.Prev"""
    try:
        document_id = data.get("document_id")
        target_type = data.get("target_type")
        target_id = data.get("target_id")

        doc = await db.documents.find_one({"id": document_id}, {"_id": 0})
        if not doc:
            raise HTTPException(status_code=404, detail="Document non trouvé")

        # Créer l'entrée d'attachment
        attachment = {
            "id": str(uuid.uuid4()),
            "document_id": document_id,
            "nom": doc.get("fichier_nom") or doc.get("titre", "Document"),
            "url": doc.get("fichier_url", ""),
            "type": doc.get("fichier_type", ""),
            "taille": doc.get("fichier_taille", 0),
            "source": "documentations",
            "added_at": datetime.now(timezone.utc).isoformat(),
            "added_by": current_user.get("id"),
            "added_by_name": f"{current_user.get('prenom', '')} {current_user.get('nom', '')}".strip()
        }

        if target_type == "work_order":
            result = await db.work_orders.update_one(
                {"id": target_id},
                {"$push": {"attachments": attachment}}
            )
            if result.modified_count == 0:
                # Essayer par _id
                from bson import ObjectId as BsonObjectId
                result = await db.work_orders.update_one(
                    {"_id": BsonObjectId(target_id)},
                    {"$push": {"attachments": attachment}}
                )
            entity_name = "l'Ordre de Travail"
        elif target_type == "improvement":
            result = await db.improvements.update_one(
                {"id": target_id},
                {"$push": {"attachments": attachment}}
            )
            entity_name = "l'Amélioration"
        elif target_type == "preventive_maintenance":
            result = await db.preventive_maintenances.update_one(
                {"id": target_id},
                {"$push": {"attachments": attachment}}
            )
            entity_name = "la Maintenance Préventive"
        else:
            raise HTTPException(status_code=400, detail="Type cible invalide")

        # Audit
        doc_name = doc.get('fichier_nom') or doc.get('titre', 'Document')
        await audit_service.log_action(
            user_id=current_user["id"],
            user_name=f"{current_user.get('prenom', '')} {current_user.get('nom', '')}",
            user_email=current_user.get("email", ""),
            action=ActionType.UPDATE,
            entity_type=EntityType.DOCUMENTATION,
            entity_id=document_id,
            entity_name=doc_name,
            details=f"a inséré \"{doc_name}\" dans {entity_name}"
        )

        return {"success": True, "message": f"Document inséré dans {entity_name}"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erreur insertion document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
